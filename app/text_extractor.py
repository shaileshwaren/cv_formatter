"""
Extract raw text from uploaded CV files (PDF or DOCX).
"""

from pathlib import Path
from PyPDF2 import PdfReader
from PyPDF2.errors import PdfReadError, EmptyFileError
from docx import Document


def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_from_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
    except (PdfReadError, EmptyFileError) as exc:
        raise ValueError("Invalid or empty PDF file.") from exc
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts)


def _extract_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)
