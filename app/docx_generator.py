"""
Generate a formatted DOCX document from structured CV data,
matching the Oxydata CV template exactly.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from app.cv_parser import CVData
from app import template_spec as T


def _has_real_skills(skills_text: str) -> bool:
    normalized = (skills_text or "").strip().lower()
    if not normalized:
        return False
    return normalized not in {"<to be filled>", "to be filled", "n/a", "na", "-"}


def _is_non_work_hobby(item: str) -> bool:
    """
    Heuristic guardrail: keep clearly personal hobbies/interests,
    exclude work/industry capability phrases that GPT may misclassify.
    """
    text = (item or "").strip()
    if not text:
        return False

    normalized = text.lower()
    if normalized in {"<to be filled>", "to be filled", "n/a", "na", "-"}:
        return False

    work_like_keywords = {
        "saas", "hr", "hiring", "employee engagement", "digital transformation",
        "sustainability", "optimization", "strategy", "strategic", "leadership",
        "stakeholder", "product", "project", "program", "consulting",
        "business development", "operations", "analytics", "data science",
        "machine learning", "ai", "automation", "cloud", "architecture",
        "software", "engineering", "enterprise", "governance", "compliance",
        "kpi", "okrs", "transformation",
    }

    # Filter out obvious work-domain phrases.
    return not any(keyword in normalized for keyword in work_like_keywords)


def _set_run(run, font_size=T.SIZE_BODY, bold=False, italic=False,
             color=T.COLOR_BLACK, font_name=T.FONT_FAMILY):
    run.bold = bold
    run.italic = italic
    run.font.size = font_size
    run.font.name = font_name
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rpr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>'))


def _set_paragraph_spacing(para, before=Pt(0), after=Pt(0), line=None):
    pf = para.paragraph_format
    pf.space_before = before
    pf.space_after = after
    if line is not None:
        pf.line_spacing = line


def _add_section_header(doc, title):
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, before=T.SPACING_BEFORE_SECTION,
                           after=T.SPACING_AFTER_SECTION_HEADER)
    run = para.add_run(title)
    _set_run(run, font_size=T.SIZE_SECTION_HEADER, bold=True,
             color=T.COLOR_DARK_BLUE)

    # Add a thin top border line above section headers
    pPr = para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="12" w:color="D9D9D9"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    return para


def _add_info_line(doc, label, value):
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, after=T.SPACING_AFTER_PERSONAL_INFO)
    run_label = para.add_run(f"{label}: ")
    _set_run(run_label, bold=True)
    run_value = para.add_run(value)
    _set_run(run_value)
    return para


def _add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    _set_paragraph_spacing(para, after=T.SPACING_AFTER_PARAGRAPH)
    for run in para.runs:
        run.clear()
    para.clear()
    run = para.add_run(text)
    _set_run(run)
    return para


def _add_skills_table(doc, skills):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for i, header_text in enumerate(["Category", "Skills"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(header_text)
        _set_run(run, bold=True)

    for skill in skills:
        row = table.add_row().cells
        for i, val in enumerate([skill.category, skill.skills]):
            row[i].text = ""
            p = row[i].paragraphs[0]
            run = p.add_run(val)
            _set_run(run)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = T.TABLE_COL_WIDTHS[0]
        row.cells[1].width = T.TABLE_COL_WIDTHS[1]

    # Light table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    return table


def _add_experience_entry(doc, entry):
    # Title — COMPANY
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, before=Pt(6), after=Pt(0))
    run_title = para.add_run(entry.title)
    _set_run(run_title, bold=True)
    run_sep = para.add_run(" \u2014 ")
    _set_run(run_sep)
    run_company = para.add_run(entry.company.upper())
    _set_run(run_company, bold=True)

    # Date range
    date_para = doc.add_paragraph()
    _set_paragraph_spacing(date_para, after=Pt(2))
    date_text = f"({entry.start_date} \u2014 {entry.end_date})"
    run_date = date_para.add_run(date_text)
    _set_run(run_date, bold=True, italic=True, color=T.COLOR_GRAY)

    # Scope line if present
    if entry.scope and entry.scope != "<To be filled>":
        scope_para = doc.add_paragraph()
        _set_paragraph_spacing(scope_para, after=Pt(2))
        run_scope = scope_para.add_run(f"Scope: {entry.scope}")
        _set_run(run_scope)

    # Bullet points
    for bullet in entry.bullets:
        _add_bullet(doc, bullet)

    # Gap between consecutive experience blocks (approx. 2 lines)
    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, after=Pt(24))


def _add_project_entry(doc, entry):
    # Company line
    _add_info_line(doc, "Company", entry.company)

    # Project line
    _add_info_line(doc, "Project", entry.project)

    # Role if present
    if entry.role and entry.role != "<To be filled>":
        _add_info_line(doc, "Project Role", entry.role)

    # Duration
    _add_info_line(doc, "Duration", entry.duration)

    # Tools
    if entry.tools and entry.tools != "<To be filled>":
        _add_info_line(doc, "Tools", entry.tools)

    # Bullets
    for bullet in entry.bullets:
        _add_bullet(doc, bullet)

    # Small spacer
    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, after=Pt(4))


def _add_education_entry(doc, entry):
    # Degree name (bold)
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, after=Pt(0))
    run = para.add_run(entry.degree)
    _set_run(run, bold=True)

    # Institution | Year | Grade
    detail_parts = [entry.institution]
    if entry.year:
        detail_parts.append(entry.year)
    if entry.grade:
        detail_parts.append(entry.grade)
    detail_para = doc.add_paragraph()
    _set_paragraph_spacing(detail_para, after=Pt(4))
    run = detail_para.add_run("  |  ".join(detail_parts))
    _set_run(run)


def _add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run(T.FOOTER_TEXT + "  |  Page ")
    _set_run(run, font_size=Pt(8), color=T.COLOR_GRAY)

    # Page number field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

    run_pg = para.add_run()
    _set_run(run_pg, font_size=Pt(8), color=T.COLOR_GRAY)
    run_pg._element.append(fld_char_begin)
    run_pg2 = para.add_run()
    _set_run(run_pg2, font_size=Pt(8), color=T.COLOR_GRAY)
    run_pg2._element.append(instr_text)
    run_pg3 = para.add_run()
    _set_run(run_pg3, font_size=Pt(8), color=T.COLOR_GRAY)
    run_pg3._element.append(fld_char_sep)
    run_pg4 = para.add_run()
    _set_run(run_pg4, font_size=Pt(8), color=T.COLOR_GRAY)
    run_pg4._element.append(fld_char_end)

    run_of = para.add_run(" of ")
    _set_run(run_of, font_size=Pt(8), color=T.COLOR_GRAY)

    # Total pages field
    fld_begin2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    fld_sep2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fld_end2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

    run_np = para.add_run()
    _set_run(run_np, font_size=Pt(8), color=T.COLOR_GRAY)
    run_np._element.append(fld_begin2)
    run_np2 = para.add_run()
    _set_run(run_np2, font_size=Pt(8), color=T.COLOR_GRAY)
    run_np2._element.append(instr2)
    run_np3 = para.add_run()
    _set_run(run_np3, font_size=Pt(8), color=T.COLOR_GRAY)
    run_np3._element.append(fld_sep2)
    run_np4 = para.add_run()
    _set_run(run_np4, font_size=Pt(8), color=T.COLOR_GRAY)
    run_np4._element.append(fld_end2)


def generate_docx(cv: CVData, output_path: str) -> str:
    doc = Document()

    # ── Page setup ──────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = T.PAGE_WIDTH
    section.page_height = T.PAGE_HEIGHT
    section.top_margin = T.MARGIN_TOP
    section.bottom_margin = T.MARGIN_BOTTOM
    section.left_margin = T.MARGIN_LEFT
    section.right_margin = T.MARGIN_RIGHT

    # ── Name ────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    _set_paragraph_spacing(name_para, after=T.SPACING_AFTER_NAME)
    run = name_para.add_run(cv.name)
    _set_run(run, font_size=T.SIZE_NAME, bold=True, color=T.COLOR_DARK_BLUE)

    # ── Personal info ───────────────────────────────────────
    _add_info_line(doc, "Nationality", cv.nationality)
    _add_info_line(doc, "Position Applied", cv.position_applied)

    exp_text = f"{cv.total_experience} Total Experience  |  {cv.relevant_experience} Relevant Experience"
    _add_info_line(doc, "Experience", exp_text)

    _add_info_line(doc, "Location", cv.location)
    _add_info_line(doc, "Notice Period", cv.notice_period)

    # ── Professional Summary ────────────────────────────────
    _add_section_header(doc, T.SECTION_PROFESSIONAL_SUMMARY)
    if cv.professional_summary:
        para = doc.add_paragraph()
        _set_paragraph_spacing(para, after=T.SPACING_AFTER_PARAGRAPH)
        run = para.add_run(cv.professional_summary)
        _set_run(run)

    # ── Technical Skills ────────────────────────────────────
    valid_technical_skills = [
        skill for skill in cv.technical_skills if _has_real_skills(skill.skills)
    ]
    if valid_technical_skills:
        _add_section_header(doc, T.SECTION_TECHNICAL_SKILLS)
        _add_skills_table(doc, valid_technical_skills)

    # ── Business Skills ─────────────────────────────────────
    if cv.business_skills:
        _add_section_header(doc, T.SECTION_BUSINESS_SKILLS)
        _add_skills_table(doc, cv.business_skills)

    # ── Soft Skills ─────────────────────────────────────────
    if cv.soft_skills:
        _add_section_header(doc, T.SECTION_SOFT_SKILLS)
        for skill in cv.soft_skills:
            _add_bullet(doc, skill)

    # ── Certifications ──────────────────────────────────────
    if cv.certifications:
        _add_section_header(doc, T.SECTION_CERTIFICATIONS)
        for cert in cv.certifications:
            _add_bullet(doc, cert)

    # ── Awards ──────────────────────────────────────────────
    if cv.awards:
        _add_section_header(doc, T.SECTION_AWARDS)
        for award in cv.awards:
            _add_bullet(doc, award)

    # ── Professional Experience ─────────────────────────────
    if cv.professional_experience:
        _add_section_header(doc, T.SECTION_PROFESSIONAL_EXPERIENCE)
        for entry in cv.professional_experience:
            _add_experience_entry(doc, entry)

    # ── Project Experience (conditional) ────────────────────
    if cv.project_experience:
        _add_section_header(doc, T.SECTION_PROJECT_EXPERIENCE)
        for entry in cv.project_experience:
            _add_project_entry(doc, entry)

    # ── Education ───────────────────────────────────────────
    if cv.education:
        _add_section_header(doc, T.SECTION_EDUCATION)
        for entry in cv.education:
            _add_education_entry(doc, entry)

    # ── Hobbies & Interests ─────────────────────────────────
    valid_hobbies = [hobby for hobby in cv.hobbies if _is_non_work_hobby(hobby)]
    if cv.languages or valid_hobbies:
        _add_section_header(doc, T.SECTION_HOBBIES)

        if cv.languages:
            lang_parts = [f"{l.language} \u2014 {l.level}" for l in cv.languages]
            lang_text = "  |  ".join(lang_parts)
            _add_info_line(doc, "Languages", lang_text)

        for hobby in valid_hobbies:
            _add_bullet(doc, hobby)

    # ── Footer ──────────────────────────────────────────────
    _add_footer(section)

    # ── Remove the initial empty paragraph ──────────────────
    if doc.paragraphs and not doc.paragraphs[0].text:
        p_element = doc.paragraphs[0]._element
        p_element.getparent().remove(p_element)

    doc.save(output_path)
    return output_path
